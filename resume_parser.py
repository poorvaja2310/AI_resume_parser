# import os
# import re
# import fitz  # PyMuPDF
# import docx
# from pathlib import Path
# from typing import Dict, Any, List, Optional
# from datetime import datetime
# import json

# try:
#     from llama_cpp import Llama
# except ImportError:
#     Llama = None


# class ResumeParser:
#     """Parse resumes (PDF/DOCX/TXT) into structured data with enhanced accuracy and optional LLM support."""

#     def __init__(self, model_path: Optional[str] = None, llm=None):
#         self.model_path = model_path
#         self.llm = llm
        
#         # Initialize LLM if model path is provided
#         if model_path and os.path.exists(model_path) and Llama is not None:
#             try:
#                 self.llm = Llama(
#                     model_path=model_path,
#                     n_ctx=2048,  # Context window
#                     n_threads=4,  # CPU threads
#                     n_gpu_layers=0,  # GPU layers (0 = CPU only)
#                     verbose=False
#                 )
#                 print(f"LLM loaded successfully from {model_path}")
#             except Exception as e:
#                 print(f"Failed to load LLM: {e}")
#                 self.llm = None

#         # Regex patterns
#         self.email_pattern = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
#         self.phone_pattern = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?\(?(?:\d{3})\)?[-.\s]?\d{3}[-.\s]?\d{4}")
#         self.linkedin_pattern = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+/?", re.I)
#         self.github_pattern = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+/?", re.I)
#         self.experience_pattern = re.compile(r'(\d+)\s*(?:years?|yrs?|\+)', re.I)

#         # Skills vocabulary
#         self.skills_vocab = {
#             'Python', 'Java', 'JavaScript', 'TypeScript', 'C', 'C++', 'C#', 'Go', 'Rust', 'PHP', 'Ruby', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Perl', 'Shell', 'Bash', 'PowerShell',
#             'HTML', 'CSS', 'SCSS', 'Sass', 'Less', 'React', 'Angular', 'Vue.js', 'Next.js', 'Nuxt.js', 'Svelte', 'jQuery', 'Bootstrap', 'Tailwind CSS', 'Material-UI', 'Chakra UI',
#             'Node.js', 'Express.js', 'Django', 'Flask', 'FastAPI', 'Spring', 'Spring Boot', 'Laravel', 'Rails', 'ASP.NET', '.NET Core', 'Gin', 'Echo',
#             'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'SQL Server', 'Cassandra', 'DynamoDB', 'Elasticsearch', 'Neo4j',
#             'AWS', 'Azure', 'GCP', 'Google Cloud', 'Docker', 'Kubernetes', 'Terraform', 'Ansible', 'Jenkins', 'GitLab CI', 'GitHub Actions', 'CircleCI', 'Travis CI',
#             'Machine Learning', 'Deep Learning', 'Data Science', 'TensorFlow', 'PyTorch', 'Keras', 'scikit-learn', 'Pandas', 'NumPy', 'Matplotlib', 'Seaborn', 'Jupyter', 'Apache Spark', 'Hadoop',
#             'Git', 'Linux', 'Unix', 'Windows', 'macOS', 'VS Code', 'IntelliJ', 'Eclipse', 'Vim', 'Emacs', 'Postman', 'Insomnia', 'Figma', 'Adobe XD', 'Photoshop', 'Illustrator'
#         }

#     # ---------------- Text Extraction ---------------- #
#     def extract_text(self, file_path: str) -> str:
#         ext = Path(file_path).suffix.lower()
#         if ext == ".pdf":
#             return self._extract_text_from_pdf(file_path)
#         elif ext in [".docx", ".doc"]:
#             return self._extract_text_from_docx(file_path)
#         elif ext == ".txt":
#             return self._extract_text_from_txt(file_path)
#         return ""

#     def _extract_text_from_pdf(self, file_path: str) -> str:
#         try:
#             doc = fitz.open(file_path)
#             text = "\n".join([page.get_text("text") for page in doc])
#             doc.close()
#             return self._clean_text(text)
#         except Exception as e:
#             print(f"PDF parsing error: {e}")
#             return ""

#     def _extract_text_from_docx(self, file_path: str) -> str:
#         try:
#             doc = docx.Document(file_path)
#             text = "\n".join([p.text for p in doc.paragraphs])
#             return self._clean_text(text)
#         except Exception as e:
#             print(f"DOCX parsing error: {e}")
#             return " "

#     def _extract_text_from_txt(self, file_path: str) -> str:
#         try:
#             with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
#                 return self._clean_text(f.read())
#         except Exception as e:
#             print(f"TXT parsing error: {e}")
#             return ""

#     def _clean_text(self, text: str) -> str:
#         text = re.sub(r'[ \t]+', ' ', text)
#         text = re.sub(r'\n\s*\n+', '\n\n', text)
#         return text.strip()

#     # ---------------- Section Extraction ---------------- #
#     def _extract_section(self, text: str, section_headers: List[str]) -> str:
#         lines = text.splitlines()
#         section_start = None
#         for i, line in enumerate(lines):
#             line_clean = re.sub(r'[^a-zA-Z\s]', '', line).strip().lower()
#             for header in section_headers:
#                 header_clean = re.sub(r'[^a-zA-Z\s]', '', header).strip().lower()
#                 if header_clean in line_clean or line_clean.startswith(header_clean):
#                     section_start = i
#                     break
#             if section_start is not None:
#                 break
#         if section_start is None:
#             return ""
#         section_content = []
#         for i in range(section_start + 1, len(lines)):
#             line = lines[i].strip()
#             if self._is_section_header(line):
#                 break
#             if line:
#                 section_content.append(line)
#         return "\n".join(section_content)

#     def _is_section_header(self, line: str) -> bool:
#         line = line.strip()
#         if not line:
#             return False
#         section_keywords = [
#             'education', 'experience', 'work experience', 'employment', 'skills', 
#             'technical skills', 'projects', 'certifications', 'achievements', 
#             'summary', 'objective', 'profile', 'contact', 'references'
#         ]
#         if line.isupper() and len(line) > 2:
#             return True
#         line_lower = line.lower()
#         return any(keyword in line_lower for keyword in section_keywords)

#     # ---------------- Parsing Methods ---------------- #
#     def parse_resume_text(self, text: str) -> Dict[str, Any]:
#         return {
#             "name": self._extract_name(text),
#             "email": self._extract_email(text),
#             "phone": self._extract_phone(text),
#             "linkedin": self._extract_linkedin(text),
#             "github": self._extract_github(text),
#             "skills": self._extract_skills(text),
#             "education": self._extract_education(text),
#             "experience": self._extract_experience(text),
#             "certifications": self._extract_certifications(text),
#             "summary": self._extract_summary(text),
#             "total_experience_years": self._extract_total_experience(text)
#         }

#     def _extract_name(self, text: str) -> str:
#         lines = [line.strip() for line in text.splitlines() if line.strip()]
#         for line in lines[:10]:
#             if self._is_likely_name_line(line):
#                 name = self._clean_name(line)
#                 if name and self._validate_name(name):
#                     return name
#         name_labels = ['name:', 'full name:', 'candidate:', 'applicant:']
#         for line in lines[:15]:
#             line_lower = line.lower()
#             for label in name_labels:
#                 if label in line_lower:
#                     name_part = line[line_lower.find(label) + len(label):].strip()
#                     name = self._clean_name(name_part)
#                     if name and self._validate_name(name):
#                         return name
#         email = self._extract_email(text)
#         if email != "Not Available":
#             prefix = email.split("@")[0]
#             name = re.sub(r'[^a-zA-Z\s]', ' ', prefix)
#             name = re.sub(r'\s+', ' ', name).strip().title()
#             if len(name.split()) >= 2:
#                 return name
#         return "Not Available"

#     def _is_likely_name_line(self, line: str) -> bool:
#         skip_patterns = [
#             r'@', r'http', r'www\.', r'\.com', r'\.org', r'\.net',
#             r'\d{3,}', r'resume', r'cv', r'curriculum',
#             r'phone', r'email', r'address', r'contact'
#         ]
#         line_lower = line.lower()
#         if any(re.search(pattern, line_lower) for pattern in skip_patterns):
#             return False
#         words = line.split()
#         if 2 <= len(words) <= 4:
#             capitalized = sum(1 for word in words if word and word[0].isupper())
#             return capitalized >= 2
#         return False

#     def _clean_name(self, name: str) -> str:
#         name = re.sub(r'\b(?:Mr\.?|Ms\.?|Mrs\.?|Dr\.?|Prof\.?)\s*', '', name, flags=re.I)
#         name = re.sub(r'\s*(?:Jr\.?|Sr\.?|III?|IV)\s*$', '', name, flags=re.I)
#         name = re.sub(r'\s+', ' ', name).strip()
#         words = []
#         for word in name.split():
#             if word.isupper() or word.islower():
#                 words.append(word.capitalize())
#             else:
#                 words.append(word)
#         return ' '.join(words)

#     def _validate_name(self, name: str) -> bool:
#         if not name or len(name) < 3:
#             return False
#         words = name.split()
#         if len(words) < 2 or len(words) > 4:
#             return False
#         job_terms = [
#             'engineer', 'developer', 'manager', 'analyst', 'designer',
#             'consultant', 'specialist', 'coordinator', 'director', 'lead',
#             'senior', 'junior', 'associate', 'intern', 'freelance'
#         ]
#         name_lower = name.lower()
#         if any(term in name_lower for term in job_terms):
#             return False
#         for word in words:
#             if len(word) < 2 or len(word) > 20:
#                 return False
#         return True

#     def _extract_email(self, text: str) -> str:
#         match = self.email_pattern.search(text)
#         return match.group(0) if match else "Not Available"

#     def _extract_phone(self, text: str) -> str:
#         match = self.phone_pattern.search(text)
#         if match:
#             phone = re.sub(r'[^\d+]', '', match.group(0))
#             if len(phone) >= 10:
#                 return phone
#         return "Not Available"

#     def _extract_linkedin(self, text: str) -> str:
#         match = self.linkedin_pattern.search(text)
#         return match.group(0) if match else "Not Available"

#     def _extract_github(self, text: str) -> str:
#         match = self.github_pattern.search(text)
#         return match.group(0) if match else "Not Available"

#     def _extract_skills(self, text: str) -> List[str]:
#         skills = set()
#         skills_section = self._extract_section(text, [
#             'skills', 'technical skills', 'core competencies', 'technologies',
#             'programming languages', 'tools', 'expertise'
#         ])
#         if skills_section:
#             skills.update(self._parse_skills_from_section(skills_section))
#         text_lower = text.lower()
#         for skill in self.skills_vocab:
#             if re.search(rf'\b{re.escape(skill.lower())}\b', text_lower):
#                 skills.add(skill)
#         # Additional patterns
#         skill_patterns = [
#             r'(?:experience (?:with|in)|proficient (?:with|in)|skilled (?:with|in)|knowledge of)\s+([A-Za-z+#.\s]+)',
#             r'(\w+(?:\.\w+)*)\s*(?:\([^)]*\))?\s*(?:programming|development|framework|library)',
#             r'(?:years? of|experience with|worked with)\s+([A-Za-z+#.\s]+)'
#         ]
#         for pattern in skill_patterns:
#             matches = re.findall(pattern, text, re.I)
#             for match in matches:
#                 for skill in re.split(r'[,;/&]', match):
#                     skill = skill.strip()
#                     if skill in self.skills_vocab:
#                         skills.add(skill)
#         return sorted(skills) if skills else ["Not Available"]

#     def _parse_skills_from_section(self, section_text: str) -> set:
#         skills = set()
#         items = re.split(r'[,;•\n\t|/]', section_text)
#         for item in items:
#             item = re.sub(r'^\W+|\W+$', '', item.strip())
#             item = re.sub(r'\s+', ' ', item)
#             if item in self.skills_vocab:
#                 skills.add(item)
#             else:
#                 item_lower = item.lower()
#                 for skill in self.skills_vocab:
#                     if skill.lower() == item_lower or skill.lower() in item_lower:
#                         skills.add(skill)
#                         break
#         return skills

#     def _extract_education(self, text: str) -> List[str]:
#         education = []
#         edu_section = self._extract_section(text, ['education', 'academic background', 'qualifications', 'degrees'])
#         if edu_section:
#             education.extend(self._parse_education_from_section(edu_section))
#         degree_patterns = [
#             r'((?:Bachelor|Master|PhD|Doctorate|Associate|B\.?[A-Za-z]*|M\.?[A-Za-z]*|Ph\.?D\.?)[^,\n]*(?:in|of)\s+[^,\n]+)',
#             r'((?:B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|MBA|Ph\.?D\.?)[^,\n]*)',
#             r'([A-Za-z\s]+(?:University|College|Institute|School)[^,\n]*)'
#         ]
#         for pattern in degree_patterns:
#             matches = re.findall(pattern, text, re.I)
#             for match in matches:
#                 cleaned = re.sub(r'\s+', ' ', match.strip())
#                 if len(cleaned) > 5 and cleaned not in education:
#                     education.append(cleaned)
#         return education if education else ["Not Available"]

#     def _parse_education_from_section(self, section_text: str) -> List[str]:
#         education = []
#         lines = [line.strip() for line in section_text.splitlines() if line.strip()]
#         current_entry = []
#         for line in lines:
#             if any(keyword in line.lower() for keyword in ['university', 'college', 'institute', 'school', 'bachelor', 'master', 'phd', 'degree']):
#                 if current_entry:
#                     education.append(' '.join(current_entry))
#                     current_entry = []
#                 current_entry.append(line)
#             elif current_entry and len(line) > 3:
#                 current_entry.append(line)
#         if current_entry:
#             education.append(' '.join(current_entry))
#         return education

#     def _extract_experience(self, text: str) -> List[str]:
#         experience = []
#         exp_section = self._extract_section(text, [
#             'experience', 'work experience', 'professional experience', 
#             'employment', 'career history', 'work history'
#         ])
#         if exp_section:
#             experience.extend(self._parse_experience_from_section(exp_section))
#         exp_patterns = [
#             r'([A-Za-z\s]+(?:Engineer|Developer|Manager|Analyst|Designer|Consultant|Specialist)[^,\n]*(?:at|@)\s+[^,\n]+)',
#             r'(\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)[^,\n]*[A-Za-z][^,\n]*)'
#         ]
#         for pattern in exp_patterns:
#             matches = re.findall(pattern, text, re.I)
#             for match in matches:
#                 cleaned = re.sub(r'\s+', ' ', match.strip())
#                 if len(cleaned) > 10 and cleaned not in experience:
#                     experience.append(cleaned)
#         return experience if experience else ["Not Available"]

#     def _parse_experience_from_section(self, section_text: str) -> List[str]:
#         experience = []
#         lines = [line.strip() for line in section_text.splitlines() if line.strip()]
#         current_entry = []
#         for line in lines:
#             if (any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'designer', 'consultant']) or
#                 re.search(r'\d{4}\s*[-–]\s*(?:\d{4}|present|current)', line.lower()) or
#                 any(keyword in line.lower() for keyword in ['company', 'corp', 'inc', 'ltd', 'llc'])):
#                 if current_entry:
#                     experience.append(' '.join(current_entry))
#                     current_entry = []
#                 current_entry.append(line)
#             elif current_entry and len(line) > 5:
#                 current_entry.append(line)
#         if current_entry:
#             experience.append(' '.join(current_entry))
#         return experience

#     def _extract_certifications(self, text: str) -> List[str]:
#         certifications = []
#         cert_section = self._extract_section(text, ['certifications', 'certificates', 'licenses', 'credentials'])
#         if cert_section:
#             lines = [line.strip() for line in cert_section.splitlines() if line.strip()]
#             for line in lines:
#                 if len(line) > 5:
#                     certifications.append(line)
#         cert_patterns = [
#             r'((?:AWS|Azure|Google|Microsoft|Oracle|Cisco|CompTIA)[^,\n]*)',
#             r'([A-Za-z\s]+Certification)'
#         ]
#         for pattern in cert_patterns:
#             matches = re.findall(pattern, text, re.I)
#             for match in matches:
#                 cleaned = re.sub(r'\s+', ' ', match.strip())
#                 if cleaned not in certifications:
#                     certifications.append(cleaned)
#         return certifications if certifications else ["Not Available"]

#     def _extract_summary(self, text: str) -> str:
#         summary = self._extract_section(text, ['summary', 'professional summary', 'career objective', 'profile', 'about me'])
#         return summary if summary else "Not Available"

#     def _extract_total_experience(self, text: str) -> str:
#         matches = self.experience_pattern.findall(text)
#         years = [int(m) for m in matches if m.isdigit()]
#         return str(max(years)) if years else "Not Available"

#     # ---------------- LLM Parsing ---------------- #
#     def _llm_parse(self, text: str) -> Optional[Dict[str, Any]]:
#         """Optional LLM enhancement using TinyLLaMA with safe JSON parsing."""
#         if not self.llm:
#             return None
            
#         try:
#             prompt = f"""
# Extract resume information as JSON with these exact keys:
# name, email, phone, linkedin, github, skills, education, experience, certifications, summary, total_experience_years

# Resume text:
# {text[:3000]}

# Return only valid JSON without any extra commentary.
# """
#             # Generate response using LLM
#             response = self.llm(
#                 prompt,
#                 max_tokens=1000,
#                 temperature=0.1,
#                 echo=False,
#                 stop=["</s>", "```"]
#             )

#             response_text = ""
#             if isinstance(response, dict) and "choices" in response:
#                 response_text = response['choices'][0].get('text', '').strip()
#             elif isinstance(response, str):
#                 response_text = response.strip()
            
#             if not response_text:
#                 return None

#             json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
#             if json_match:
#                 json_str = json_match.group(0)
#                 try:
#                     return json.loads(json_str)
#                 except json.JSONDecodeError:
#                     cleaned_json = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_str)
#                     try:
#                         return json.loads(cleaned_json)
#                     except json.JSONDecodeError:
#                         print("Warning: LLM returned invalid JSON, skipping LLM enhancement")
#                         return None
#             else:
#                 print("Warning: LLM output does not contain JSON, skipping LLM enhancement")
#                 return None

#         except Exception as e:
#             print(f"LLM parsing error: {e}")
#             return None

#     # ---------------- Public Parsing ---------------- #
#     def parse_resume(self, file_path: str) -> Dict[str, Any]:
#         text = self.extract_text(file_path)
#         parsed = self.parse_resume_text(text)
#         if self.llm:
#             llm_result = self._llm_parse(text)
#             if llm_result:
#                 parsed.update({k: v for k, v in llm_result.items() if v and v != "Not Available"})
#         return parsed

import os
import re
import fitz  # PyMuPDF
import docx
from pathlib import Path
from typing import Dict, Any, List, Optional
from datetime import datetime
import json

try:
    from llama_cpp import Llama
except ImportError:
    Llama = None


class ResumeParser:
    """Parse resumes (PDF/DOCX/TXT) into structured data with enhanced accuracy and optional LLM support."""

    def __init__(self, model_path: Optional[str] = None, llm=None):
        self.model_path = model_path
        self.llm = llm
        
        # Initialize LLM if model path is provided
        if model_path and os.path.exists(model_path) and Llama is not None:
            try:
                self.llm = Llama(
                    model_path=model_path,
                    n_ctx=2048,  # Context window
                    n_threads=4,  # CPU threads
                    n_gpu_layers=0,  # GPU layers (0 = CPU only)
                    verbose=False
                )
                print(f"LLM loaded successfully from {model_path}")
            except Exception as e:
                print(f"Failed to load LLM: {e}")
                self.llm = None

        # Regex patterns
        self.email_pattern = re.compile(r"\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b")
        self.phone_pattern = re.compile(r"(?:\+?\d{1,3}[-.\s]?)?\(?(?:\d{3})\)?[-.\s]?\d{3}[-.\s]?\d{4}")
        self.linkedin_pattern = re.compile(r"(?:https?://)?(?:www\.)?linkedin\.com/in/[\w\-]+/?", re.I)
        self.github_pattern = re.compile(r"(?:https?://)?(?:www\.)?github\.com/[\w\-]+/?", re.I)
        self.experience_pattern = re.compile(r'(\d+)\s*(?:years?|yrs?|\+)', re.I)

        # Skills vocabulary
        self.skills_vocab = {
            'Python', 'Java', 'JavaScript', 'TypeScript', 'C', 'C++', 'C#', 'Go', 'Rust', 'PHP', 'Ruby', 'Swift', 'Kotlin', 'Scala', 'R', 'MATLAB', 'Perl', 'Shell', 'Bash', 'PowerShell',
            'HTML', 'CSS', 'SCSS', 'Sass', 'Less', 'React', 'Angular', 'Vue.js', 'Next.js', 'Nuxt.js', 'Svelte', 'jQuery', 'Bootstrap', 'Tailwind CSS', 'Material-UI', 'Chakra UI',
            'Node.js', 'Express.js', 'Django', 'Flask', 'FastAPI', 'Spring', 'Spring Boot', 'Laravel', 'Rails', 'ASP.NET', '.NET Core', 'Gin', 'Echo',
            'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'SQLite', 'Oracle', 'SQL Server', 'Cassandra', 'DynamoDB', 'Elasticsearch', 'Neo4j',
            'AWS', 'Azure', 'GCP', 'Google Cloud', 'Docker', 'Kubernetes', 'Terraform', 'Ansible', 'Jenkins', 'GitLab CI', 'GitHub Actions', 'CircleCI', 'Travis CI',
            'Machine Learning', 'Deep Learning', 'Data Science', 'TensorFlow', 'PyTorch', 'Keras', 'scikit-learn', 'Pandas', 'NumPy', 'Matplotlib', 'Seaborn', 'Jupyter', 'Apache Spark', 'Hadoop',
            'Git', 'Linux', 'Unix', 'Windows', 'macOS', 'VS Code', 'IntelliJ', 'Eclipse', 'Vim', 'Emacs', 'Postman', 'Insomnia', 'Figma', 'Adobe XD', 'Photoshop', 'Illustrator'
        }

    # ---------------- Text Extraction ---------------- #
    def extract_text(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext == ".pdf":
            return self._extract_text_from_pdf(file_path)
        elif ext in [".docx", ".doc"]:
            return self._extract_text_from_docx(file_path)
        elif ext == ".txt":
            return self._extract_text_from_txt(file_path)
        return ""

    def _extract_text_from_pdf(self, file_path: str) -> str:
        try:
            doc = fitz.open(file_path)
            text = "\n".join([page.get_text("text") for page in doc])
            doc.close()
            return self._clean_text(text)
        except Exception as e:
            print(f"PDF parsing error: {e}")
            return ""

    def _extract_text_from_docx(self, file_path: str) -> str:
        try:
            doc = docx.Document(file_path)
            text = "\n".join([p.text for p in doc.paragraphs])
            return self._clean_text(text)
        except Exception as e:
            print(f"DOCX parsing error: {e}")
            return " "

    def _extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                return self._clean_text(f.read())
        except Exception as e:
            print(f"TXT parsing error: {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        text = re.sub(r'[ \t]+', ' ', text)
        text = re.sub(r'\n\s*\n+', '\n\n', text)
        return text.strip()

    # ---------------- Section Extraction ---------------- #
    def _extract_section(self, text: str, section_headers: List[str]) -> str:
        lines = text.splitlines()
        section_start = None
        for i, line in enumerate(lines):
            line_clean = re.sub(r'[^a-zA-Z\s]', '', line).strip().lower()
            for header in section_headers:
                header_clean = re.sub(r'[^a-zA-Z\s]', '', header).strip().lower()
                if header_clean in line_clean or line_clean.startswith(header_clean):
                    section_start = i
                    break
            if section_start is not None:
                break
        if section_start is None:
            return ""
        section_content = []
        for i in range(section_start + 1, len(lines)):
            line = lines[i].strip()
            if self._is_section_header(line):
                break
            if line:
                section_content.append(line)
        return "\n".join(section_content)

    def _is_section_header(self, line: str) -> bool:
        line = line.strip()
        if not line:
            return False
        section_keywords = [
            'education', 'experience', 'work experience', 'employment', 'skills', 
            'technical skills', 'projects', 'certifications', 'achievements', 
            'summary', 'objective', 'profile', 'contact', 'references'
        ]
        if line.isupper() and len(line) > 2:
            return True
        line_lower = line.lower()
        return any(keyword in line_lower for keyword in section_keywords)

    # ---------------- Parsing Methods ---------------- #
    def parse_resume_text(self, text: str) -> Dict[str, Any]:
        return {
            "name": self._extract_name(text),
            "email": self._extract_email(text),
            "phone": self._extract_phone(text),
            "linkedin": self._extract_linkedin(text),
            "github": self._extract_github(text),
            "skills": self._extract_skills(text),
            "education": self._extract_education(text),
            "experience": self._extract_experience(text),
            "certifications": self._extract_certifications(text),
            "summary": self._extract_summary(text),
            "total_experience_years": self._extract_total_experience(text)
        }

    def _extract_name(self, text: str) -> str:
        lines = [line.strip() for line in text.splitlines() if line.strip()]
        for line in lines[:10]:
            if self._is_likely_name_line(line):
                name = self._clean_name(line)
                if name and self._validate_name(name):
                    return name
        name_labels = ['name:', 'full name:', 'candidate:', 'applicant:']
        for line in lines[:15]:
            line_lower = line.lower()
            for label in name_labels:
                if label in line_lower:
                    name_part = line[line_lower.find(label) + len(label):].strip()
                    name = self._clean_name(name_part)
                    if name and self._validate_name(name):
                        return name
        email = self._extract_email(text)
        if email != "Not Available":
            prefix = email.split("@")[0]
            name = re.sub(r'[^a-zA-Z\s]', ' ', prefix)
            name = re.sub(r'\s+', ' ', name).strip().title()
            if len(name.split()) >= 2:
                return name
        return "Not Available"

    def _is_likely_name_line(self, line: str) -> bool:
        skip_patterns = [
            r'@', r'http', r'www\.', r'\.com', r'\.org', r'\.net',
            r'\d{3,}', r'resume', r'cv', r'curriculum',
            r'phone', r'email', r'address', r'contact'
        ]
        line_lower = line.lower()
        if any(re.search(pattern, line_lower) for pattern in skip_patterns):
            return False
        words = line.split()
        if 2 <= len(words) <= 4:
            capitalized = sum(1 for word in words if word and word[0].isupper())
            return capitalized >= 2
        return False

    def _clean_name(self, name: str) -> str:
        name = re.sub(r'\b(?:Mr\.?|Ms\.?|Mrs\.?|Dr\.?|Prof\.?)\s*', '', name, flags=re.I)
        name = re.sub(r'\s*(?:Jr\.?|Sr\.?|III?|IV)\s*$', '', name, flags=re.I)
        name = re.sub(r'\s+', ' ', name).strip()
        words = []
        for word in name.split():
            if word.isupper() or word.islower():
                words.append(word.capitalize())
            else:
                words.append(word)
        return ' '.join(words)

    def _validate_name(self, name: str) -> bool:
        if not name or len(name) < 3:
            return False
        words = name.split()
        if len(words) < 2 or len(words) > 4:
            return False
        job_terms = [
            'engineer', 'developer', 'manager', 'analyst', 'designer',
            'consultant', 'specialist', 'coordinator', 'director', 'lead',
            'senior', 'junior', 'associate', 'intern', 'freelance'
        ]
        name_lower = name.lower()
        if any(term in name_lower for term in job_terms):
            return False
        for word in words:
            if len(word) < 2 or len(word) > 20:
                return False
        return True

    def _extract_email(self, text: str) -> str:
        match = self.email_pattern.search(text)
        return match.group(0) if match else "Not Available"

    def _extract_phone(self, text: str) -> str:
        match = self.phone_pattern.search(text)
        if match:
            phone = re.sub(r'[^\d+]', '', match.group(0))
            if len(phone) >= 10:
                return phone
        return "Not Available"

    def _extract_linkedin(self, text: str) -> str:
        match = self.linkedin_pattern.search(text)
        return match.group(0) if match else "Not Available"

    def _extract_github(self, text: str) -> str:
        match = self.github_pattern.search(text)
        return match.group(0) if match else "Not Available"

    def _extract_skills(self, text: str) -> List[str]:
        skills = set()
        skills_section = self._extract_section(text, [
            'skills', 'technical skills', 'core competencies', 'technologies',
            'programming languages', 'tools', 'expertise'
        ])
        if skills_section:
            skills.update(self._parse_skills_from_section(skills_section))
        text_lower = text.lower()
        for skill in self.skills_vocab:
            if re.search(rf'\b{re.escape(skill.lower())}\b', text_lower):
                skills.add(skill)
        # Additional patterns
        skill_patterns = [
            r'(?:experience (?:with|in)|proficient (?:with|in)|skilled (?:with|in)|knowledge of)\s+([A-Za-z+#.\s]+)',
            r'(\w+(?:\.\w+)*)\s*(?:\([^)]*\))?\s*(?:programming|development|framework|library)',
            r'(?:years? of|experience with|worked with)\s+([A-Za-z+#.\s]+)'
        ]
        for pattern in skill_patterns:
            matches = re.findall(pattern, text, re.I)
            for match in matches:
                for skill in re.split(r'[,;/&]', match):
                    skill = skill.strip()
                    if skill in self.skills_vocab:
                        skills.add(skill)
        return sorted(skills) if skills else ["Not Available"]

    def _parse_skills_from_section(self, section_text: str) -> set:
        skills = set()
        items = re.split(r'[,;•\n\t|/]', section_text)
        for item in items:
            item = re.sub(r'^\W+|\W+$', '', item.strip())
            item = re.sub(r'\s+', ' ', item)
            if item in self.skills_vocab:
                skills.add(item)
            else:
                item_lower = item.lower()
                for skill in self.skills_vocab:
                    if skill.lower() == item_lower or skill.lower() in item_lower:
                        skills.add(skill)
                        break
        return skills

    def _extract_education(self, text: str) -> List[str]:
        education = []
        edu_section = self._extract_section(text, ['education', 'academic background', 'qualifications', 'degrees'])
        if edu_section:
            education.extend(self._parse_education_from_section(edu_section))
        degree_patterns = [
            r'((?:Bachelor|Master|PhD|Doctorate|Associate|B\.?[A-Za-z]*|M\.?[A-Za-z]*|Ph\.?D\.?)[^,\n]*(?:in|of)\s+[^,\n]+)',
            r'((?:B\.?S\.?|M\.?S\.?|B\.?A\.?|M\.?A\.?|MBA|Ph\.?D\.?)[^,\n]*)',
            r'([A-Za-z\s]+(?:University|College|Institute|School)[^,\n]*)'
        ]
        for pattern in degree_patterns:
            matches = re.findall(pattern, text, re.I)
            for match in matches:
                cleaned = re.sub(r'\s+', ' ', match.strip())
                if len(cleaned) > 5 and cleaned not in education:
                    education.append(cleaned)
        return education if education else ["Not Available"]

    def _parse_education_from_section(self, section_text: str) -> List[str]:
        education = []
        lines = [line.strip() for line in section_text.splitlines() if line.strip()]
        current_entry = []
        for line in lines:
            if any(keyword in line.lower() for keyword in ['university', 'college', 'institute', 'school', 'bachelor', 'master', 'phd', 'degree']):
                if current_entry:
                    education.append(' '.join(current_entry))
                    current_entry = []
                current_entry.append(line)
            elif current_entry and len(line) > 3:
                current_entry.append(line)
        if current_entry:
            education.append(' '.join(current_entry))
        return education

    def _extract_experience(self, text: str) -> List[str]:
        experience = []
        exp_section = self._extract_section(text, [
            'experience', 'work experience', 'professional experience', 
            'employment', 'career history', 'work history'
        ])
        if exp_section:
            experience.extend(self._parse_experience_from_section(exp_section))
        exp_patterns = [
            r'([A-Za-z\s]+(?:Engineer|Developer|Manager|Analyst|Designer|Consultant|Specialist)[^,\n]*(?:at|@)\s+[^,\n]+)',
            r'(\d{4}\s*[-–]\s*(?:\d{4}|Present|Current)[^,\n]*[A-Za-z][^,\n]*)'
        ]
        for pattern in exp_patterns:
            matches = re.findall(pattern, text, re.I)
            for match in matches:
                cleaned = re.sub(r'\s+', ' ', match.strip())
                if len(cleaned) > 10 and cleaned not in experience:
                    experience.append(cleaned)
        return experience if experience else ["Not Available"]

    def _parse_experience_from_section(self, section_text: str) -> List[str]:
        experience = []
        lines = [line.strip() for line in section_text.splitlines() if line.strip()]
        current_entry = []
        for line in lines:
            if (any(keyword in line.lower() for keyword in ['engineer', 'developer', 'manager', 'analyst', 'designer', 'consultant']) or
                re.search(r'\d{4}\s*[-–]\s*(?:\d{4}|present|current)', line.lower()) or
                any(keyword in line.lower() for keyword in ['company', 'corp', 'inc', 'ltd', 'llc'])):
                if current_entry:
                    experience.append(' '.join(current_entry))
                    current_entry = []
                current_entry.append(line)
            elif current_entry and len(line) > 5:
                current_entry.append(line)
        if current_entry:
            experience.append(' '.join(current_entry))
        return experience

    def _extract_certifications(self, text: str) -> List[str]:
        certifications = []
        cert_section = self._extract_section(text, ['certifications', 'certificates', 'licenses', 'credentials'])
        if cert_section:
            lines = [line.strip() for line in cert_section.splitlines() if line.strip()]
            for line in lines:
                if len(line) > 5:
                    certifications.append(line)
        cert_patterns = [
            r'((?:AWS|Azure|Google|Microsoft|Oracle|Cisco|CompTIA)[^,\n]*)',
            r'([A-Za-z\s]+Certification)'
        ]
        for pattern in cert_patterns:
            matches = re.findall(pattern, text, re.I)
            for match in matches:
                cleaned = re.sub(r'\s+', ' ', match.strip())
                if cleaned not in certifications:
                    certifications.append(cleaned)
        return certifications if certifications else ["Not Available"]

    def _extract_summary(self, text: str) -> str:
        summary = self._extract_section(text, ['summary', 'professional summary', 'career objective', 'profile', 'about me'])
        return summary if summary else "Not Available"

    def _extract_total_experience(self, text: str) -> str:
        matches = self.experience_pattern.findall(text)
        years = [int(m) for m in matches if m.isdigit()]
        return str(max(years)) if years else "Not Available"

    # ---------------- LLM Parsing ---------------- #
    def _llm_parse(self, text: str) -> Optional[Dict[str, Any]]:
        """Optional LLM enhancement using TinyLLaMA with safe JSON parsing."""
        if not self.llm:
            return None
            
        try:
            prompt = f"""
Extract resume information as JSON with these exact keys:
name, email, phone, linkedin, github, skills, education, experience, certifications, summary, total_experience_years

Resume text:
{text[:3000]}

Return only valid JSON without any extra commentary.
"""
            # Generate response using LLM
            response = self.llm(
                prompt,
                max_tokens=1000,
                temperature=0.1,
                echo=False,
                stop=["</s>", "```"]
            )

            response_text = ""
            if isinstance(response, dict) and "choices" in response:
                response_text = response['choices'][0].get('text', '').strip()
            elif isinstance(response, str):
                response_text = response.strip()
            
            if not response_text:
                return None

            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group(0)
                try:
                    return json.loads(json_str)
                except json.JSONDecodeError:
                    cleaned_json = re.sub(r'[\x00-\x1f\x7f-\x9f]', '', json_str)
                    try:
                        return json.loads(cleaned_json)
                    except json.JSONDecodeError:
                        print("Warning: LLM returned invalid JSON, skipping LLM enhancement")
                        return None
            else:
                print("Warning: LLM output does not contain JSON, skipping LLM enhancement")
                return None

        except Exception as e:
            print(f"LLM parsing error: {e}")
            return None

    # ---------------- Public Parsing ---------------- #
    def parse_resume(self, file_path: str) -> Dict[str, Any]:
        text = self.extract_text(file_path)
        parsed = self.parse_resume_text(text)
        if self.llm:
            llm_result = self._llm_parse(text)
            if llm_result:
                parsed.update({k: v for k, v in llm_result.items() if v and v != "Not Available"})
        return parsed